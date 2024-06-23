import io
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENTATION
from talon.models import Branch, Talon


def create_docx_with_tables(talons, from_date, to_date):
    # Create a new Word document
    doc = Document()

    # Set the orientation to landscape
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height

    title = doc.add_paragraph(
        f'{from_date.strftime("%Y-%m-%d")} дан {to_date.strftime("%Y-%m-%d")} гача ҳолатига жамият корхона ва '
        'ташкилотларда меҳнат муҳофазаси бўйича олинган огоҳлантириш талонлари тўғрисида\n\n'
        'МТУ лар бўйича олинган талонлар'
    )
    run = title.runs[0]
    run.bold = True
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    # Add the first table title
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = p.add_run("1-чи жадвал")
    run.font.size = Pt(14)

    # Create the first table
    table1 = doc.add_table(rows=3, cols=18)
    table1.style = 'Table Grid'

    # Define the header data
    header_data = [
        ("Тошкент РЖУ", 3), ("Қўқон РЖУ", 3), ("Бухоро РЖУ", 3),
        ("Қарши РЖУ", 3), ("Термиз РЖУ", 3), ("Қунгирот РЖУ", 3)
    ]

    # Add headers
    hdr_cells = table1.rows[0].cells
    col_index = 0
    for header, span in header_data:
        cell = hdr_cells[col_index]
        cell.text = header
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for i in range(1, span):
            cell.merge(hdr_cells[col_index + i])
        col_index += span

    # Define sub-headers
    sub_headers = ["№1", "№2", "№3"]
    subhdr_cells = table1.rows[1].cells
    col_index = 0
    for _ in range(6):
        for sub_header in sub_headers:
            subhdr_cells[col_index].text = sub_header
            subhdr_cells[col_index].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            col_index += 1

    data_row = [97, 33, "-", 139, 64, "-", 172, 45, "-", 137, 27, "-", 88, 32, 1, 118, 23, "-"]
    row_cells = table1.add_row().cells
    for col_index, item in enumerate(data_row):
        row_cells[col_index].text = str(item)
        row_cells[col_index].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add the final row of the first table
    final_row = [130, 163, 217, 164, 121, 141]
    final_row_cells = table1.add_row().cells
    for col_index, item in enumerate(final_row):
        final_row_cells[col_index * 3].text = str(item)
        final_row_cells[col_index * 3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for i in range(1, 3):
            final_row_cells[col_index * 3 + i].merge(final_row_cells[col_index * 3])

    # Add a blank paragraph between tables
    title2 = doc.add_paragraph('МТУ  лар бўйича январ - май ойлари учун олинган талонларнинг ушлаб қолинган суммалари')
    run = title2.runs[0]
    run.bold = True
    title2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title2.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    # Define the data for the second table
    table2_data = [
        ["Тошкент РЖУ", "Қўқон РЖУ", "Бухоро РЖУ", "Қарши РЖУ", "Термиз РЖУ", "Қунгирот РЖУ"],
        [29627445, 10141988, 65275881, 87001607, 32907607, 12168867],
        [237123395]
    ]

    # Add the second table title
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = p.add_run("2-чи жадвал")
    run.font.size = Pt(14)

    # Add the second table
    table2 = doc.add_table(rows=2, cols=6)
    table2.style = 'Table Grid'
    hdr_cells = table2.rows[0].cells
    for i, hdr in enumerate(table2_data[0]):
        hdr_cells[i].text = hdr
        hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add the data row
    data_cells = table2.rows[1].cells
    for i, data in enumerate(table2_data[1]):
        data_cells[i].text = str(data)
        data_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add the total row spanning all columns
    row = table2.add_row()
    cell = row.cells[0]
    cell.text = 'ЖАМИ'
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cell.merge(row.cells[4])
    row.cells[5].text = str(table2_data[2][0])
    row.cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add a blank paragraph between tables
    doc.add_paragraph()

    # Add the third table title
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = p.add_run("3-чи жадвал")
    run.font.size = Pt(14)

    branches = Branch.objects.all()
    for branch in branches:
        branch_title = doc.add_paragraph(branch.name)
        run = branch_title.runs[0]
        run.bold = True
        branch_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        branch_title.runs[0].font.color.rgb = RGBColor(0, 0, 0)

        organizations = branch.branch_organizations.all()
        if organizations.exists():
            for organization in organizations:
                organization_title = doc.add_paragraph(organization.name)
                run = organization_title.runs[0]
                run.bold = True
                organization_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                organization_title.runs[0].font.color.rgb = RGBColor(0, 0, 0)

                # Add a new table for this organization
                table3 = doc.add_table(rows=1, cols=9)
                table3.style = 'Table Grid'

                # Define table headers
                hdr_cells = table3.rows[0].cells
                hdr_cells[0].text = 'Т/р'
                hdr_cells[1].text = 'Ф.И.Ш.'
                hdr_cells[2].text = 'Лавозими'
                hdr_cells[3].text = 'Олинган санаси'
                hdr_cells[4].text = 'Талон рақами'
                hdr_cells[5].text = 'Олинган сабаби'
                hdr_cells[6].text = 'Интизомий жазо буйруқ рақами, санаси ва тури'
                hdr_cells[7].text = 'Интизомий жазо оқибатида олиб қолинган пул миқдори'
                hdr_cells[8].text = 'Изоҳ'

                talons = Talon.objects.filter(organization=organization)
                # Populate table with data
                for idx, talon in enumerate(talons, 1):
                    row_cells = table3.add_row().cells
                    row_cells[0].text = str(idx)
                    row_cells[1].text = talon.fullname
                    row_cells[2].text = talon.position
                    row_cells[3].text = talon.date_received.strftime("%d.%m.%Y")
                    row_cells[4].text = talon.talon_number
                    row_cells[5].text = talon.reason_received
                    row_cells[6].text = f"Рақам: {talon.discipline_order}, Сана: {talon.discipline_order_date.strftime('%d.%m.%Y')}, Тури: {talon.discipline_type}"
                    row_cells[7].text = f"{talon.consequence_amount:.2f}" if talon.consequence_amount else ""
                    row_cells[8].text = talon.note or ""
                doc.add_paragraph('Жами: ')
        else:
            # Add a new table for this branch
            table3 = doc.add_table(rows=1, cols=9)
            table3.style = 'Table Grid'

            # Define table headers
            hdr_cells = table3.rows[0].cells
            hdr_cells[0].text = 'Т/р'
            hdr_cells[1].text = 'Ф.И.Ш.'
            hdr_cells[2].text = 'Лавозими'
            hdr_cells[3].text = 'Олинган санаси'
            hdr_cells[4].text = 'Талон рақами'
            hdr_cells[5].text = 'Олинган сабаби'
            hdr_cells[6].text = 'Интизомий жазо буйруқ рақами, санаси ва тури'
            hdr_cells[7].text = 'Интизомий жазо оқибатида олиб қолинган пул миқдори'
            hdr_cells[8].text = 'Изоҳ'

            talons = Talon.objects.filter(branch=branch, organization__isnull=True)
            # Populate table with data
            for idx, talon in enumerate(talons, 1):
                row_cells = table3.add_row().cells
                row_cells[0].text = str(idx)
                row_cells[1].text = talon.fullname
                row_cells[2].text = talon.position
                row_cells[3].text = talon.date_received.strftime("%d.%m.%Y")
                row_cells[4].text = talon.talon_number
                row_cells[5].text = talon.reason_received
                row_cells[6].text = f"Рақам: {talon.discipline_order}, Сана: {talon.discipline_order_date.strftime('%d.%m.%Y')}, Тури: {talon.discipline_type}"
                row_cells[7].text = f"{talon.consequence_amount:.2f}" if talon.consequence_amount else ""
                row_cells[8].text = talon.note or ""
            doc.add_paragraph('Жами: ')
        all_branch_amount = doc.add_paragraph(f'{branch.name} - Жами:')
        run = all_branch_amount.runs[0]
        run.bold = True
        all_branch_amount.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        all_branch_amount.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    all_amount = doc.add_paragraph(f'Умумий жами:  ')
    run = all_amount.runs[0]
    run.bold = True
    all_amount.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    all_amount.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    # Save the document to a BytesIO buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer

